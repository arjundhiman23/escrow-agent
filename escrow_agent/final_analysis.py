"""Build the Final Analysis document.

Summarises the KB (sanction letter + sanction note/CAM + escrow agreement),
cross-checks the CATRA/TRA outputs against the Agreement clauses, and issues a
per-item and overall verdict: OVERBREACH / UNDERBREACH / COMPLIANT.

Breach semantics used throughout the agent:
  OVERBREACH  — utilisation/outflow above what the clause permits (cap exceeded,
                unauthorised withdrawal, payment out of Order of Priority,
                distribution made while senior obligations unfunded).
  UNDERBREACH — funding/obligation below what the clause requires (reserve
                shortfall, unfunded senior priority, deposit not routed to escrow).
On the initial inputs (no bank statement yet) the verdicts are on PROJECTED basis
from the CAM; the same checks re-run on actuals when the statement is ingested.
"""
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x7A, 0x33)
AMBER = RGBColor(0xB8, 0x86, 0x00)


def _style(doc):
    s = doc.styles["Normal"]
    s.font.name = "Arial"
    s.font.size = Pt(10)


def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Arial"
        r.font.color.rgb = DARK
    return h


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def _status_color(t, col):
    for row in t.rows[1:]:
        cell = row.cells[col]
        txt = cell.text.strip()
        if txt == "COMPLIANT":
            color = GREEN
        elif txt in ("UNDERBREACH", "OVERBREACH"):
            color = RED
        else:  # MARGINAL, N-A, PENDING ACTUALS, ATTENTION
            color = AMBER
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = color
                r.font.bold = True


def build_final_analysis(extract, reserve_checks, out_path="output/Final_Analysis.docx"):
    doc = Document()
    _style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("FINAL ANALYSIS — ESCROW ACCOUNT COMPLIANCE\nKanpur Lucknow Expressway Private Limited")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = DARK
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Generated from Knowledge Base (Sanction Letter, Credit Approval Memo, Escrow Agreement) — {date.today().strftime('%d-%m-%Y')}")
    r.font.size = Pt(9); r.font.italic = True

    # ---------------- 1. Executive summary & verdict ----------------
    _h(doc, "1. Executive Summary and Verdict")
    fails = [c for c in reserve_checks if c["status"] in ("UNDERBREACH", "OVERBREACH")]
    marginal = [c for c in reserve_checks if c["status"] == "MARGINAL"]
    over = [c for c in fails if c["status"] == "OVERBREACH"]
    under = [c for c in fails if c["status"] == "UNDERBREACH"]
    if over:
        verdict, vcol = "OVERBREACH", RED
    elif under:
        verdict, vcol = "UNDERBREACH", RED
    else:
        verdict, vcol = "COMPLIANT", GREEN
    p = doc.add_paragraph()
    p.add_run("Overall account status (projected basis): ").font.size = Pt(11)
    vr = p.add_run(verdict)
    vr.font.size = Pt(13); vr.font.bold = True; vr.font.color.rgb = vcol
    doc.add_paragraph(
        f"Basis: {len(reserve_checks)} covenant checks computed from the CAM base-case projections against the "
        f"Escrow Agreement requirements — {len(fails)} fail ({len(under)} UNDERBREACH, {len(over)} OVERBREACH), "
        f"{len(marginal)} are marginal (within proxy tolerance, analyst to verify against the actual debt-service "
        "schedule), and the rest are compliant or not applicable in the year concerned. No actual bank statement "
        "has been ingested yet — once the escrow account statement is provided, every check below re-runs on "
        "actuals and the verdict is restated on an actuals basis. UNDERBREACH means a reserve or senior obligation "
        "is funded below the level the Agreement requires; OVERBREACH would mean utilisation above a permitted cap "
        "or a payment made out of the Order of Priority."
    )

    # ---------------- 2. Deal summary ----------------
    _h(doc, "2. Deal Summary (from KB)")
    _table(doc, ["Item", "Detail"], [
        ["Borrower", "Kanpur Lucknow Expressway Private Limited (SPV of PNC Infratech Limited)"],
        ["Facility", "Rupee Term Loan of Rs. 779.75 crore — Axis Bank (sole lender, Escrow Bank, Lenders' Representative)"],
        ["Project", "Six-lane (upgradable to eight) Kanpur Lucknow Expressway incl. Spur, UP — HAM, Bharatmala (Pkg-I)"],
        ["Sanction Letter", "AXISB/LC/North/2022-23/2133 dated 26-08-2022; cash-flow waterfall at clause 20 (p.33-34)"],
        ["Escrow Agreement", "e-stamp 27-09-2022; deposits cl. 2.3(A)(a)-(g); Order of Priority cl. 2.3(B)(a)-(m)"],
        ["Annuity structure", "60% of BPC in 30 biannual instalments over 15 years from 180th day of COD; 40% construction grant in 10 x 4% instalments; termination payment 65% of annuity (Rs. 610.22 cr)"],
        ["Note", "Agreement PDF is misnamed 'Kallagam Meensuruti' — content verified as the KLEPL–Axis Bank agreement"],
    ], widths=[1.6, 5.4])

    # ---------------- 3. CATRA summary ----------------
    _h(doc, "3. CATRA Classification Framework (generated)")
    doc.add_paragraph(
        "Debit-side categories follow the Escrow Agreement Order of Priority (authoritative); the Sanction Letter "
        "clause-20 waterfall maps onto it step-for-step. Credit-side categories follow the permitted deposits in "
        "clause 2.3(A). Full detail, keyword rules and clause references are in CATRA_Master.xlsx; the transaction "
        "classifier now runs against this taxonomy (config/categories.yaml)."
    )
    _table(doc, ["Priority", "Code", "Category", "Agreement", "Sanction cl.20"],
           [[c["priority"], c["code"], c["name"], f"2.3(B)({c['item']})",
             ", ".join(str(s["step"]) for s in extract["sanction_waterfall"] if s["maps_to"] == c["code"]) or "—"]
            for c in extract["order_of_priority"]],
           widths=[0.7, 1.1, 2.6, 1.1, 1.2])
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Waterfall divergences identified: ").font.bold = True
    for d in extract["waterfall_divergences"]:
        doc.add_paragraph(d, style="List Bullet")

    # ---------------- 4. TRA summary ----------------
    _h(doc, "4. TRA Analysis (from Sanction Note / CAM Projected P&L)")
    pnl = extract["projected_pnl"]
    doc.add_paragraph(
        "Projected inflows are entirely annuity-linked (TPC annuity, interest on annuity at reducing balance, and "
        "O&M receipts paid with each instalment). Projected outflows are O&M, MMR provisioning, facility interest "
        "and scheduled principal. Full FY26–FY35 detail, the CATRA-mapped cashflow view and the semi-annual profile "
        "are in TRA_Analysis.xlsx."
    )
    fy_show = ["FY26", "FY27", "FY30", "FY33", "FY35"]
    idx = [pnl["fye"].index(f) for f in fy_show]
    bs = extract["projected_balance_sheet"]
    _table(doc, ["Rs. crore"] + fy_show, [
        ["Total income"] + [pnl["income_total"][i] for i in idx],
        ["EBITDA"] + [pnl["ebitda"][i] for i in idx],
        ["Interest (facility)"] + [pnl["interest"][i] for i in idx],
        ["Principal (CMLTD)"] + [bs["cmltd"][i] for i in idx],
        ["PAT"] + [pnl["pat"][i] for i in idx],
        ["DSRA fund"] + [bs["dsra_fund"][i] for i in idx],
        ["WCR fund"] + [bs["working_capital_reserve"][i] for i in idx],
    ], widths=[1.8, 1.04, 1.04, 1.04, 1.04, 1.04])

    # ---------------- 5. Clause cross-check ----------------
    _h(doc, "5. Cross-Check Against Escrow Agreement Clauses")
    rows = []
    # aggregate reserve checks by covenant: worst year
    for cvid, label in [("CV2", "DSRA = 2 quarters debt service (2.3(B)(j))"),
                        ("CV1", "WCR = 6 months interest at COD (2.3(B)(i))"),
                        ("CV3", "MMRA per base case (2.3(B)(k))")]:
        sub = [c for c in reserve_checks if c["covenant"].startswith(cvid)]
        bad = [c for c in sub if c["status"] in ("UNDERBREACH", "OVERBREACH")]
        marg = [c for c in sub if c["status"] == "MARGINAL"]
        if bad:
            worst = min(bad, key=lambda c: c["gap"])
            extra = f"; {len(marg)} further year(s) marginal" if marg else ""
            rows.append([label, worst["status"],
                         f"Worst year {worst['fy']}: required {worst['required']} vs provided {worst['provided']} (gap {worst['gap']}){extra}. {worst['note']}"])
        elif marg:
            worst = min(marg, key=lambda c: c["gap"])
            rows.append([label, "MARGINAL",
                         f"{len(marg)} year(s) within proxy tolerance (worst {worst['fy']}: gap {worst['gap']}). {worst['note']}"])
        else:
            rows.append([label, "COMPLIANT", "Projected fund meets requirement in all applicable years"])
    rows += [
        ["Order of Priority routing (2.3(B) chapeau; CV6)", "PENDING ACTUALS",
         "Requires the escrow bank statement: every withdrawal must route via the Sub-Accounts in priority order on Monthly Cash Transfer Dates"],
        ["Annuity deposit within 1 business day (2.3(A)(c); CV5)", "PENDING ACTUALS",
         "Requires statement credit dates vs NHAI annuity payment dates"],
        ["50% surplus prepayment on each annuity (Sanction; CV4)", "PENDING ACTUALS",
         "Requires actual surplus computation per annuity cycle and prepayment evidence"],
        ["Bonus 100% to prepayment only if DSRA & MMR created (CV8)", "PENDING ACTUALS",
         "Conditional; check at bonus receipt"],
        ["Shortfall priority on Debt Service (2.11(iii); CV7)", "PENDING ACTUALS",
         "Applies only if a debt-service shortfall event occurs"],
        ["Subordinate Debt sub-account (2.3(B)(h))", "ATTENTION",
         "Present in Agreement, absent from Sanction cl.20 waterfall — confirm treatment with lender before classifying any sub-debt payment"],
    ]
    t = _table(doc, ["Clause / Covenant", "Status", "Finding"], rows, widths=[2.3, 1.2, 3.5])
    _status_color(t, 1)

    # ---------------- 6. Detailed reserve table ----------------
    _h(doc, "6. Reserve Adequacy Detail (projected basis)")
    t2 = _table(doc, ["Covenant", "FY", "Required", "Provided", "Gap", "Status"],
                [[c["covenant"], c["fy"], c["required"], c["provided"], c["gap"], c["status"]]
                 for c in reserve_checks],
                widths=[2.4, 0.7, 1.0, 1.0, 0.9, 1.0])
    _status_color(t2, 5)

    # ---------------- 7. Next inputs ----------------
    _h(doc, "7. Inputs Required to Move from Projected to Actuals Basis")
    for item in [
        "Escrow account bank statement (xlsx/csv/pdf) — enables classification against the generated CATRA, waterfall-order validation, and actual-vs-TRA variance",
        "Supplementary Escrow Agreement — referenced by the Sanction Letter for the surplus waterfall; needed to finalise CV4 mechanics",
        "COD date / annuity payment schedule as invoiced — to pin the semi-annual TRA profile to actual instalment dates",
        "Confirmation of WCR sizing interpretation (6 months interest vs CAM base-case level) and of Subordinate Debt treatment",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(out_path)
    return out_path, verdict
