"""Word report generation (BRD §6.6, §8.2) — a formatted management report
for credit/compliance review and audit sign-off, generated with python-docx."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .models import Taxonomy, Transaction

QUARTERS = ("Q1", "Q2", "Q3", "Q4")
NAVY = RGBColor(0x1F, 0x4E, 0x79)


def _fmt(x) -> str:
    return f"{x:,.2f}" if x is not None else "—"


def _style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for lvl in (1, 2):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Arial"
        h.font.color.rgb = NAVY


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = str(v)
    return t


def generate(out_path: str | Path, txns: list[Transaction], taxonomy: Taxonomy,
             agg: dict, var: dict, wf_exceptions: list[dict],
             manual_conds: list[dict], sanction: dict, warnings: list[str],
             audit) -> Path:
    doc = Document()
    _style(doc)

    title = doc.add_heading("Escrow Account Transaction Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Deal: {sanction.get('deal_name', 'N/A')}    |    "
                            f"Generated: {datetime.now():%d-%m-%Y %H:%M}    |    "
                            f"Scope: Main escrow account only (BRD §4.1)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    # 1. Executive summary -------------------------------------------------
    doc.add_heading("1. Executive Summary", level=1)
    fy_dr = sum(v["Dr"] for v in agg["totals"].values())
    fy_cr = sum(v["Cr"] for v in agg["totals"].values())
    n_material = sum(r["material"] for r in var["rows"])
    n_review = sum(t.needs_review for t in txns)
    n_conflict = sum(t.is_escalated_conflict for t in txns)
    doc.add_paragraph(
        f"The agent processed {len(txns)} transactions on the main escrow account. "
        f"Full-year credits total INR {_fmt(fy_cr)} against debits of INR {_fmt(fy_dr)}. "
        f"{sum(1 for t in txns if t.category_code)} transactions were classified against the "
        f"TRA/CATRA taxonomy ({sum(t.classification_source == 'RULE' for t in txns)} by rule, "
        f"{sum(t.classification_source == 'AI' for t in txns)} by AI); {n_review} entries are pending "
        f"in the analyst review queue and {n_conflict} entries were escalated for manual override due to "
        f"potentially inaccurate narrations. Waterfall and sanction-condition validation flagged "
        f"{len(wf_exceptions)} exceptions ({sum(1 for e in wf_exceptions if e['side'] == 'Debit')} debit-side, "
        f"{sum(1 for e in wf_exceptions if e['side'] == 'Credit')} credit-side). Variance analysis identified "
        f"{n_material} material deviations from Sanction Note projections. "
        f"{sum(t.is_internal_transfer for t in txns)} internal company transfers were flagged separately, "
        f"and {sum(t.is_duplicate for t in txns)} duplicate entries were excluded from aggregates.")

    if warnings:
        doc.add_paragraph("Data quality notes from ingestion:").runs[0].bold = True
        for w in warnings[:15]:
            doc.add_paragraph(w, style="List Bullet")
        if len(warnings) > 15:
            doc.add_paragraph(f"...and {len(warnings) - 15} further notes — see the Excel audit trail.",
                              style="List Bullet")

    # 2. Quarter-wise summary ----------------------------------------------
    doc.add_heading("2. Quarter-wise Summary", level=1)
    rows = [["Total Credits (Inflows)"] + [_fmt(agg["totals"][q]["Cr"]) for q in QUARTERS] + [_fmt(fy_cr)],
            ["Total Debits (Outflows)"] + [_fmt(agg["totals"][q]["Dr"]) for q in QUARTERS] + [_fmt(fy_dr)],
            ["Net Movement"] + [_fmt(agg["totals"][q]["Cr"] - agg["totals"][q]["Dr"]) for q in QUARTERS]
            + [_fmt(fy_cr - fy_dr)],
            ["of which Internal Transfers (Cr)"] + [_fmt(agg["internal_transfers"][q]["Cr"]) for q in QUARTERS]
            + [_fmt(sum(agg["internal_transfers"][q]["Cr"] for q in QUARTERS))],
            ["of which Internal Transfers (Dr)"] + [_fmt(agg["internal_transfers"][q]["Dr"]) for q in QUARTERS]
            + [_fmt(sum(agg["internal_transfers"][q]["Dr"] for q in QUARTERS))]]
    _table(doc, ["(INR)", "Q1", "Q2", "Q3", "Q4", "Full Year"], rows)
    doc.add_paragraph("Category-level quarter-wise breakdowns are provided in the accompanying "
                      "Excel workbook ('Quarter Summary' sheet).").runs[0].italic = True

    # 3. Waterfall compliance -----------------------------------------------
    doc.add_heading("3. Waterfall & Sanction Condition Compliance", level=1)
    doc.add_paragraph(
        "Actual transactions were validated against the priority-of-payments structure and "
        "conditions extracted from the Sanction Letter. Debit-side and credit-side transactions "
        "were validated independently; each violation carries a specific reason code.")
    if wf_exceptions:
        _table(doc, ["Side", "Violation Type", "Date", "Category", "Amount (INR)", "Commentary"],
               [[e["side"], e["violation_type"], e["date"].strftime("%d-%m-%Y"),
                 e["category"], _fmt(e["amount"]), e["detail"]] for e in wf_exceptions])
    else:
        doc.add_paragraph("No waterfall or condition violations were detected in the period analysed.")
    if manual_conds:
        doc.add_paragraph("Conditions requiring manual verification (not machine-checkable):").runs[0].bold = True
        for c in manual_conds:
            doc.add_paragraph(f"{c.get('id', '')}: {c.get('description', '')}", style="List Bullet")

    # 4. Variance analysis ---------------------------------------------------
    doc.add_heading("4. Actual vs Projected Variance Analysis", level=1)
    material = [r for r in var["rows"] if r["material"]]
    doc.add_paragraph(
        f"Actual category flows were compared with the quarter-wise projections in the Sanction Note. "
        f"{len(material)} of {len(var['rows'])} category-quarter combinations breached the configured "
        f"materiality thresholds. Material deviations:")
    if material:
        _table(doc, ["Qtr", "Side", "Category", "Actual", "Projected", "Var (INR)", "Var %", "Remarks"],
               [[r["quarter"], r["side"], r["name"], _fmt(r["actual"]), _fmt(r["projected"]),
                 _fmt(r["var_abs"]),
                 f'{r["var_pct"]:.1f}%' if r["var_pct"] is not None else "n/a",
                 r["remarks"]] for r in material])
    else:
        doc.add_paragraph("No material deviations identified.")

    # 5. Consolidated exceptions summary -------------------------------------
    doc.add_heading("5. Consolidated Exceptions Summary", level=1)
    doc.add_paragraph("Most material deviations per quarter, for management review and audit sign-off (BRD §6.5):")
    any_summary = False
    for q in QUARTERS:
        top = var["exceptions_summary"].get(q, [])
        if not top:
            continue
        any_summary = True
        p = doc.add_paragraph()
        p.add_run(f"{q}: ").bold = True
        parts = []
        for r in top:
            direction = ("unbudgeted" if r["projected"] == 0 else
                         "above plan" if r["var_abs"] > 0 else "below plan")
            parts.append(f"{r['name']} ({r['side'].lower()}) {direction} by INR {_fmt(abs(r['var_abs']))}")
        p.add_run("; ".join(parts) + ".")
    if not any_summary:
        doc.add_paragraph("No material deviations in any quarter.")

    # 6. Review queue ---------------------------------------------------------
    doc.add_heading("6. Analyst Review Queue & Escalations", level=1)
    pending = [t for t in txns if t.needs_review or t.is_escalated_conflict]
    if pending:
        doc.add_paragraph(
            f"{len(pending)} transactions require analyst attention before the analysis can be treated as final. "
            f"Escalated items (narration conflicts with transaction type) require manual override; ordinary "
            f"review items require category confirmation. Details are on the 'Review Queue' sheet of the Excel workbook.")
        _table(doc, ["Queue", "Date", "Type", "Amount (INR)", "Narration", "Reason"],
               [[("ESCALATED" if t.is_escalated_conflict else "REVIEW"),
                 t.txn_date.strftime("%d-%m-%Y"), t.txn_type, _fmt(t.amount),
                 (t.narration or "(blank)")[:60],
                 (t.conflict_reason or t.review_reason)[:90]] for t in pending[:25]])
    else:
        doc.add_paragraph("The review queue is empty — all transactions were classified with confidence.")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    audit.log("REPORT_GENERATION", f"Word report written: {out_path.name}")
    return out_path
