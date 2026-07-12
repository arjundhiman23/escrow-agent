"""Sanction document analysis (BRD §6.4).

Two supported inputs:
  1. A pre-extracted / analyst-reviewed JSON ("sanction extract") — see schema
     below and sample_data/sample_sanction_extract.json.
  2. A raw Sanction Letter / Sanction Note in PDF or Word: the module extracts
     text (OCR escalation per Risk R-02 is flagged, not silently attempted) and
     uses the Anthropic API to produce the same JSON schema, which SHOULD then
     be reviewed by an analyst before the run is treated as authoritative.

Sanction extract JSON schema
----------------------------
{
  "deal_name": str,
  "waterfall": [                       # priority of payments, 1 = highest
    {"priority": 1, "category_code": "STAT_PAY", "description": "..."} , ...
  ],
  "permitted_inflow_sources": ["OTH_REV", "PROC_EQUITY", ...],
  "conditions": [                      # machine-checkable sanction conditions
    {"id": "C1", "type": "surplus_requires_prior", "params": {"required_code":
       "DEBT_SERV", "target_code": "SURPLUS_DIST"},
     "description": "Surplus may be distributed only after debt servicing in the same cycle"},
    {"id": "C2", "type": "category_cap_per_quarter", "params": {"category_code":
       "INTERCO", "cap": 5000000}, "description": "Inter-company transfers capped at 50 lakh/quarter"},
    {"id": "C3", "type": "manual", "description": "..."}   # unverifiable-from-data -> listed for manual check
  ],
  "projections": {                     # from the Sanction Note, per quarter
    "Q1": {"inflows": {"OTH_REV": 100.0, ...}, "outflows": {"DEBT_SERV": 50.0, ...}},
    "Q2": {...}, "Q3": {...}, "Q4": {...}
  }
}
Per-quarter figures are mandatory (annual totals alone are insufficient — BRD §6.4).
"""
from __future__ import annotations

import json
from pathlib import Path

REQUIRED_KEYS = ("waterfall", "permitted_inflow_sources", "conditions", "projections")


def load_sanction_extract(path: str | Path, taxonomy, audit=None) -> dict:
    with open(path, encoding="utf-8") as fh:
        data = json.load(fh)
    problems = validate_extract(data, taxonomy)
    if problems:
        raise ValueError("Sanction extract validation failed:\n- " + "\n- ".join(problems))
    if audit:
        n_proj_q = sum(1 for q in ("Q1", "Q2", "Q3", "Q4") if q in data.get("projections", {}))
        audit.log("SANCTION_ANALYSIS",
                  f"Sanction extract loaded: {len(data['waterfall'])} waterfall levels, "
                  f"{len(data['conditions'])} conditions, projections for {n_proj_q} quarters.")
    return data


def validate_extract(data: dict, taxonomy) -> list[str]:
    problems = []
    for key in REQUIRED_KEYS:
        if key not in data:
            problems.append(f"missing key '{key}'")
    if problems:
        return problems
    valid = set(taxonomy.by_code())
    for level in data["waterfall"]:
        if level.get("category_code") not in valid:
            problems.append(f"waterfall references unknown category '{level.get('category_code')}'")
    for code in data["permitted_inflow_sources"]:
        if code not in valid:
            problems.append(f"permitted_inflow_sources references unknown category '{code}'")
    proj = data["projections"]
    quarters_present = [q for q in ("Q1", "Q2", "Q3", "Q4") if q in proj]
    if not quarters_present:
        problems.append("projections must be quarter-by-quarter (Q1..Q4); annual totals alone are insufficient (BRD §6.4)")
    for q in quarters_present:
        for side in ("inflows", "outflows"):
            for code in proj[q].get(side, {}):
                if code not in valid:
                    problems.append(f"projection {q}/{side} references unknown category '{code}'")
    return problems


# ---------------------------------------------------------------------------
# AI extraction from raw Sanction Letter / Note (PDF or Word)
# ---------------------------------------------------------------------------
def extract_from_document(path: str | Path, taxonomy, settings: dict,
                          audit=None, usage=None) -> dict:
    """Extract the sanction JSON from a raw PDF/DOCX using the Anthropic API.
    The result is written next to the source as <name>.extract.json for
    analyst review; review before treating a run as authoritative."""
    path = Path(path)
    text = _document_text(path)
    if len(text.strip()) < 200:
        raise ValueError(
            f"{path.name}: little or no machine-readable text found. Likely a "
            "scanned PDF — apply OCR pre-processing or provide a reviewed "
            "sanction extract JSON (Risk R-02).")

    import anthropic
    client = anthropic.Anthropic()
    model = settings["classification"].get("ai_model_extraction", "claude-sonnet-5")
    taxonomy_text = "\n".join(f"- {c.code} ({c.side}): {c.name}" for c in taxonomy.all())

    prompt = f"""Extract structured data from this Sanction Letter / Sanction Note for an escrow account.

Category taxonomy (use ONLY these codes):
{taxonomy_text}

Return ONLY JSON with this exact schema:
{{"deal_name": str,
  "waterfall": [{{"priority": int, "category_code": str, "description": str}}],
  "permitted_inflow_sources": [str],
  "conditions": [{{"id": str, "type": "surplus_requires_prior"|"category_cap_per_quarter"|"manual",
                   "params": {{}}, "description": str}}],
  "projections": {{"Q1": {{"inflows": {{code: number}}, "outflows": {{code: number}}}},
                   "Q2": {{...}}, "Q3": {{...}}, "Q4": {{...}}}}}}

Rules: projections MUST be per quarter by category (never annual-only); amounts in INR;
conditions that cannot be verified from transaction data get type "manual".

Document text:
{text[:150000]}"""
    resp = client.messages.create(model=model, max_tokens=8000,
                                  messages=[{"role": "user", "content": prompt}])
    if usage is not None and getattr(resp, "usage", None):
        usage.add(model, resp.usage.input_tokens, resp.usage.output_tokens)
    raw = "".join(b.text for b in resp.content if b.type == "text")
    raw = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    data = json.loads(raw)

    problems = validate_extract(data, taxonomy)
    if problems:
        raise ValueError("AI sanction extraction produced invalid data — manual "
                         "review required:\n- " + "\n- ".join(problems))
    out = path.with_suffix(path.suffix + ".extract.json")
    out.write_text(json.dumps(data, indent=2), encoding="utf-8")
    if audit:
        audit.log("SANCTION_ANALYSIS",
                  f"AI-extracted sanction structure from {path.name} -> {out.name} "
                  f"(review recommended before sign-off).")
    return data


def _document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    if suffix in (".docx", ".doc"):
        import docx as _docx
        d = _docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    raise ValueError(f"Unsupported sanction document format: {suffix}")
